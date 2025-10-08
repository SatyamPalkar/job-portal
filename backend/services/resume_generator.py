"""Resume generation service for creating formatted resume documents."""
import json
from typing import Dict, Optional
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER


class ResumeGenerator:
    """Generate formatted resume documents in various formats."""
    
    def __init__(self, output_dir: str = "./generated_resumes"):
        """Initialize the generator."""
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_pdf(self, resume_data: Dict, filename: Optional[str] = None) -> str:
        """
        Generate a PDF resume.
        
        Args:
            resume_data: Dictionary containing resume content
            filename: Optional custom filename
            
        Returns:
            Path to generated PDF file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.pdf"
        
        filepath = self.output_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Container for PDF elements
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(0, 0, 128),
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=RGBColor(0, 0, 128),
            spaceAfter=6,
            spaceBefore=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6
        )
        
        # Add name/title
        if 'name' in resume_data or 'title' in resume_data:
            name = resume_data.get('name', 'Your Name')
            story.append(Paragraph(name, title_style))
            story.append(Spacer(1, 12))
        
        # Add contact info
        if 'email' in resume_data or 'phone' in resume_data:
            contact = []
            if resume_data.get('email'):
                contact.append(resume_data['email'])
            if resume_data.get('phone'):
                contact.append(resume_data['phone'])
            if resume_data.get('location'):
                contact.append(resume_data['location'])
            
            contact_text = " | ".join(contact)
            story.append(Paragraph(contact_text, body_style))
            story.append(Spacer(1, 12))
        
        # Add summary
        if resume_data.get('summary'):
            story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
            story.append(Paragraph(resume_data['summary'], body_style))
            story.append(Spacer(1, 12))
        
        # Add experience
        if resume_data.get('experience'):
            story.append(Paragraph("WORK EXPERIENCE", heading_style))
            experiences = self._parse_json_field(resume_data['experience'])
            for exp in experiences:
                if isinstance(exp, dict):
                    title = exp.get('title', '')
                    story.append(Paragraph(f"<b>{title}</b>", body_style))
                    for desc in exp.get('description', []):
                        story.append(Paragraph(f"• {desc}", body_style))
            story.append(Spacer(1, 12))
        
        # Add education
        if resume_data.get('education'):
            story.append(Paragraph("EDUCATION", heading_style))
            educations = self._parse_json_field(resume_data['education'])
            for edu in educations:
                if isinstance(edu, dict):
                    title = edu.get('title', '')
                    story.append(Paragraph(f"<b>{title}</b>", body_style))
                    for desc in edu.get('description', []):
                        story.append(Paragraph(desc, body_style))
            story.append(Spacer(1, 12))
        
        # Add skills
        if resume_data.get('skills'):
            story.append(Paragraph("SKILLS", heading_style))
            skills = self._parse_json_field(resume_data['skills'])
            if isinstance(skills, list):
                skills_text = " • ".join(skills)
                story.append(Paragraph(skills_text, body_style))
            story.append(Spacer(1, 12))
        
        # Add projects
        if resume_data.get('projects'):
            story.append(Paragraph("PROJECTS", heading_style))
            projects = self._parse_json_field(resume_data['projects'])
            for proj in projects:
                if isinstance(proj, dict):
                    title = proj.get('title', '')
                    story.append(Paragraph(f"<b>{title}</b>", body_style))
                    for desc in proj.get('description', []):
                        story.append(Paragraph(f"• {desc}", body_style))
            story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        
        return str(filepath)
    
    def generate_docx(self, resume_data: Dict, filename: Optional[str] = None) -> str:
        """
        Generate a DOCX resume.
        
        Args:
            resume_data: Dictionary containing resume content
            filename: Optional custom filename
            
        Returns:
            Path to generated DOCX file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.docx"
        
        filepath = self.output_dir / filename
        
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add name/title
        if 'name' in resume_data or 'title' in resume_data:
            name = resume_data.get('name', 'Your Name')
            heading = doc.add_heading(name, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 128)
        
        # Add contact info
        if 'email' in resume_data or 'phone' in resume_data:
            contact = []
            if resume_data.get('email'):
                contact.append(resume_data['email'])
            if resume_data.get('phone'):
                contact.append(resume_data['phone'])
            if resume_data.get('location'):
                contact.append(resume_data['location'])
            
            contact_para = doc.add_paragraph(" | ".join(contact))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Add summary
        if resume_data.get('summary'):
            doc.add_heading('PROFESSIONAL SUMMARY', level=2)
            doc.add_paragraph(resume_data['summary'])
        
        # Add experience
        if resume_data.get('experience'):
            doc.add_heading('WORK EXPERIENCE', level=2)
            experiences = self._parse_json_field(resume_data['experience'])
            for exp in experiences:
                if isinstance(exp, dict):
                    title = exp.get('title', '')
                    p = doc.add_paragraph()
                    run = p.add_run(title)
                    run.bold = True
                    
                    for desc in exp.get('description', []):
                        doc.add_paragraph(desc, style='List Bullet')
        
        # Add education
        if resume_data.get('education'):
            doc.add_heading('EDUCATION', level=2)
            educations = self._parse_json_field(resume_data['education'])
            for edu in educations:
                if isinstance(edu, dict):
                    title = edu.get('title', '')
                    p = doc.add_paragraph()
                    run = p.add_run(title)
                    run.bold = True
                    
                    for desc in edu.get('description', []):
                        doc.add_paragraph(desc)
        
        # Add skills
        if resume_data.get('skills'):
            doc.add_heading('SKILLS', level=2)
            skills = self._parse_json_field(resume_data['skills'])
            if isinstance(skills, list):
                doc.add_paragraph(" • ".join(skills))
            else:
                doc.add_paragraph(str(skills))
        
        # Add projects
        if resume_data.get('projects'):
            doc.add_heading('PROJECTS', level=2)
            projects = self._parse_json_field(resume_data['projects'])
            for proj in projects:
                if isinstance(proj, dict):
                    title = proj.get('title', '')
                    p = doc.add_paragraph()
                    run = p.add_run(title)
                    run.bold = True
                    
                    for desc in proj.get('description', []):
                        doc.add_paragraph(desc, style='List Bullet')
        
        # Save document
        doc.save(str(filepath))
        
        return str(filepath)
    
    def generate_txt(self, resume_data: Dict, filename: Optional[str] = None) -> str:
        """
        Generate a plain text resume.
        
        Args:
            resume_data: Dictionary containing resume content
            filename: Optional custom filename
            
        Returns:
            Path to generated TXT file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.txt"
        
        filepath = self.output_dir / filename
        
        lines = []
        
        # Add name/title
        if 'name' in resume_data:
            lines.append(resume_data['name'].upper())
            lines.append("=" * len(resume_data['name']))
            lines.append("")
        
        # Add contact info
        if 'email' in resume_data or 'phone' in resume_data:
            contact = []
            if resume_data.get('email'):
                contact.append(f"Email: {resume_data['email']}")
            if resume_data.get('phone'):
                contact.append(f"Phone: {resume_data['phone']}")
            if resume_data.get('location'):
                contact.append(f"Location: {resume_data['location']}")
            
            lines.extend(contact)
            lines.append("")
        
        # Add summary
        if resume_data.get('summary'):
            lines.append("PROFESSIONAL SUMMARY")
            lines.append("-" * 20)
            lines.append(resume_data['summary'])
            lines.append("")
        
        # Add experience
        if resume_data.get('experience'):
            lines.append("WORK EXPERIENCE")
            lines.append("-" * 15)
            experiences = self._parse_json_field(resume_data['experience'])
            for exp in experiences:
                if isinstance(exp, dict):
                    lines.append(exp.get('title', ''))
                    for desc in exp.get('description', []):
                        lines.append(f"  • {desc}")
                    lines.append("")
        
        # Add education
        if resume_data.get('education'):
            lines.append("EDUCATION")
            lines.append("-" * 9)
            educations = self._parse_json_field(resume_data['education'])
            for edu in educations:
                if isinstance(edu, dict):
                    lines.append(edu.get('title', ''))
                    for desc in edu.get('description', []):
                        lines.append(f"  {desc}")
                    lines.append("")
        
        # Add skills
        if resume_data.get('skills'):
            lines.append("SKILLS")
            lines.append("-" * 6)
            skills = self._parse_json_field(resume_data['skills'])
            if isinstance(skills, list):
                lines.append(", ".join(skills))
            else:
                lines.append(str(skills))
            lines.append("")
        
        # Write to file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
        
        return str(filepath)
    
    def _parse_json_field(self, field):
        """Parse a field that might be JSON string or already parsed."""
        if isinstance(field, str):
            try:
                return json.loads(field)
            except:
                return field
        return field


