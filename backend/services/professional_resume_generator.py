"""Professional Resume Generator with modern formatting."""
import json
from typing import Dict, Optional, List
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib import colors


class ProfessionalResumeGenerator:
    """Generate professional, ATS-friendly resume documents."""
    
    def __init__(self, output_dir: str = "./generated_resumes"):
        """Initialize the generator."""
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_pdf(self, resume_data: Dict, filename: Optional[str] = None) -> str:
        """Generate a professional PDF resume with modern formatting."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.pdf"
        
        filepath = self.output_dir / filename
        
        # Create PDF with A4 size
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.6*inch,
            bottomMargin=0.6*inch
        )
        
        story = []
        styles = getSampleStyleSheet()
        
        # Define professional styles
        name_style = ParagraphStyle(
            'Name',
            parent=styles['Heading1'],
            fontSize=22,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        contact_style = ParagraphStyle(
            'Contact',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#555555'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica'
        )
        
        section_heading_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=colors.HexColor('#2c5aa0'),
            spaceBefore=14,
            spaceAfter=8,
            fontName='Helvetica-Bold',
            borderWidth=0,
            borderPadding=0,
            borderColor=colors.HexColor('#2c5aa0'),
            underlineWidth=2,
        )
        
        job_title_style = ParagraphStyle(
            'JobTitle',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=2,
            fontName='Helvetica-Bold'
        )
        
        company_style = ParagraphStyle(
            'Company',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#555555'),
            spaceAfter=4,
            fontName='Helvetica-Oblique'
        )
        
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#333333'),
            spaceAfter=4,
            fontName='Helvetica',
            leading=14
        )
        
        bullet_style = ParagraphStyle(
            'Bullet',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#333333'),
            spaceAfter=3,
            leftIndent=20,
            fontName='Helvetica',
            leading=13
        )
        
        # Header - Name
        name = resume_data.get('name', resume_data.get('full_name', 'Your Name'))
        story.append(Paragraph(f"<b>{name.upper()}</b>", name_style))
        
        # Contact Information
        contact_parts = []
        if resume_data.get('email'):
            contact_parts.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_parts.append(resume_data['phone'])
        if resume_data.get('location'):
            contact_parts.append(resume_data['location'])
        if resume_data.get('linkedin'):
            contact_parts.append(resume_data['linkedin'])
        
        if contact_parts:
            contact_text = " • ".join(contact_parts)
            story.append(Paragraph(contact_text, contact_style))
        
        # Horizontal line
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#2c5aa0'), spaceBefore=0, spaceAfter=12))
        
        # Professional Summary
        if resume_data.get('summary'):
            story.append(Paragraph("<b>PROFESSIONAL SUMMARY</b>", section_heading_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#2c5aa0'), spaceBefore=2, spaceAfter=6))
            summary_text = resume_data['summary']
            story.append(Paragraph(summary_text, body_style))
            story.append(Spacer(1, 8))
        
        # Skills
        if resume_data.get('skills'):
            story.append(Paragraph("<b>SKILLS</b>", section_heading_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#2c5aa0'), spaceBefore=2, spaceAfter=6))
            
            skills = self._parse_json_field(resume_data['skills'])
            if isinstance(skills, list):
                # Group skills into rows of 3
                skill_text = " • ".join(skills)
                story.append(Paragraph(skill_text, body_style))
            else:
                story.append(Paragraph(str(skills), body_style))
            story.append(Spacer(1, 8))
        
        # Professional Experience
        if resume_data.get('experience'):
            story.append(Paragraph("<b>PROFESSIONAL EXPERIENCE</b>", section_heading_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#2c5aa0'), spaceBefore=2, spaceAfter=6))
            
            experiences = self._parse_json_field(resume_data['experience'])
            for exp in experiences:
                if isinstance(exp, dict):
                    # Job title and company
                    title_text = exp.get('title', 'Position')
                    story.append(Paragraph(f"<b>{title_text}</b>", job_title_style))
                    
                    # Company and dates
                    company_info = exp.get('company', '')
                    if not company_info and 'description' in exp and exp['description']:
                        # Sometimes title contains company info
                        company_info = title_text
                    
                    story.append(Paragraph(company_info, company_style))
                    
                    # Responsibilities/achievements
                    descriptions = exp.get('description', [])
                    if isinstance(descriptions, str):
                        descriptions = [descriptions]
                    
                    for desc in descriptions:
                        if desc.strip():
                            # Add bullet point
                            story.append(Paragraph(f"• {desc.strip()}", bullet_style))
                    
                    story.append(Spacer(1, 8))
        
        # Education
        if resume_data.get('education'):
            story.append(Paragraph("<b>EDUCATION</b>", section_heading_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#2c5aa0'), spaceBefore=2, spaceAfter=6))
            
            educations = self._parse_json_field(resume_data['education'])
            for edu in educations:
                if isinstance(edu, dict):
                    degree = edu.get('title', 'Degree')
                    story.append(Paragraph(f"<b>{degree}</b>", job_title_style))
                    
                    descriptions = edu.get('description', [])
                    if isinstance(descriptions, str):
                        descriptions = [descriptions]
                    
                    for desc in descriptions:
                        if desc.strip():
                            story.append(Paragraph(desc.strip(), body_style))
                    
                    story.append(Spacer(1, 6))
        
        # Projects
        if resume_data.get('projects'):
            story.append(Paragraph("<b>PROJECTS</b>", section_heading_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#2c5aa0'), spaceBefore=2, spaceAfter=6))
            
            projects = self._parse_json_field(resume_data['projects'])
            for proj in projects:
                if isinstance(proj, dict):
                    proj_name = proj.get('title', 'Project')
                    story.append(Paragraph(f"<b>{proj_name}</b>", job_title_style))
                    
                    descriptions = proj.get('description', [])
                    if isinstance(descriptions, str):
                        descriptions = [descriptions]
                    
                    for desc in descriptions:
                        if desc.strip():
                            story.append(Paragraph(f"• {desc.strip()}", bullet_style))
                    
                    story.append(Spacer(1, 6))
        
        # Certifications
        if resume_data.get('certifications'):
            story.append(Paragraph("<b>CERTIFICATIONS</b>", section_heading_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#2c5aa0'), spaceBefore=2, spaceAfter=6))
            
            certs = self._parse_json_field(resume_data['certifications'])
            for cert in certs:
                if isinstance(cert, dict):
                    cert_name = cert.get('title', 'Certification')
                    story.append(Paragraph(f"• {cert_name}", bullet_style))
                elif isinstance(cert, str):
                    story.append(Paragraph(f"• {cert}", bullet_style))
        
        # Build PDF
        doc.build(story)
        return str(filepath)
    
    def generate_docx(self, resume_data: Dict, filename: Optional[str] = None) -> str:
        """Generate a professional DOCX resume."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.docx"
        
        filepath = self.output_dir / filename
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Name - Large and bold
        name = resume_data.get('name', resume_data.get('full_name', 'Your Name'))
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name.upper())
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.color.rgb = RGBColor(26, 26, 26)
        
        # Contact information
        contact_parts = []
        if resume_data.get('email'):
            contact_parts.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_parts.append(resume_data['phone'])
        if resume_data.get('location'):
            contact_parts.append(resume_data['location'])
        
        if contact_parts:
            contact_para = doc.add_paragraph(" • ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.runs[0].font.size = Pt(10)
            contact_para.runs[0].font.color.rgb = RGBColor(85, 85, 85)
        
        doc.add_paragraph()  # Spacing
        
        # Professional Summary
        if resume_data.get('summary'):
            self._add_section_heading(doc, 'PROFESSIONAL SUMMARY')
            summary_para = doc.add_paragraph(resume_data['summary'])
            summary_para.runs[0].font.size = Pt(10)
            doc.add_paragraph()  # Spacing
        
        # Skills
        if resume_data.get('skills'):
            self._add_section_heading(doc, 'SKILLS')
            skills = self._parse_json_field(resume_data['skills'])
            if isinstance(skills, list):
                skills_text = " • ".join(skills)
                skills_para = doc.add_paragraph(skills_text)
                skills_para.runs[0].font.size = Pt(10)
            doc.add_paragraph()  # Spacing
        
        # Professional Experience
        if resume_data.get('experience'):
            self._add_section_heading(doc, 'PROFESSIONAL EXPERIENCE')
            experiences = self._parse_json_field(resume_data['experience'])
            
            for exp in experiences:
                if isinstance(exp, dict):
                    # Job title
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(exp.get('title', 'Position'))
                    title_run.bold = True
                    title_run.font.size = Pt(11)
                    
                    # Company
                    company_para = doc.add_paragraph(exp.get('company', ''))
                    company_para.runs[0].italic = True
                    company_para.runs[0].font.size = Pt(10)
                    company_para.runs[0].font.color.rgb = RGBColor(85, 85, 85)
                    
                    # Descriptions
                    descriptions = exp.get('description', [])
                    if isinstance(descriptions, str):
                        descriptions = [descriptions]
                    
                    for desc in descriptions:
                        if desc.strip():
                            bullet_para = doc.add_paragraph(desc.strip(), style='List Bullet')
                            bullet_para.runs[0].font.size = Pt(10)
                    
                    doc.add_paragraph()  # Spacing between jobs
        
        # Education
        if resume_data.get('education'):
            self._add_section_heading(doc, 'EDUCATION')
            educations = self._parse_json_field(resume_data['education'])
            
            for edu in educations:
                if isinstance(edu, dict):
                    degree_para = doc.add_paragraph()
                    degree_run = degree_para.add_run(edu.get('title', 'Degree'))
                    degree_run.bold = True
                    degree_run.font.size = Pt(11)
                    
                    descriptions = edu.get('description', [])
                    if isinstance(descriptions, str):
                        descriptions = [descriptions]
                    
                    for desc in descriptions:
                        if desc.strip():
                            desc_para = doc.add_paragraph(desc.strip())
                            desc_para.runs[0].font.size = Pt(10)
        
        # Projects
        if resume_data.get('projects'):
            self._add_section_heading(doc, 'PROJECTS')
            projects = self._parse_json_field(resume_data['projects'])
            
            for proj in projects:
                if isinstance(proj, dict):
                    proj_para = doc.add_paragraph()
                    proj_run = proj_para.add_run(proj.get('title', 'Project'))
                    proj_run.bold = True
                    proj_run.font.size = Pt(11)
                    
                    descriptions = proj.get('description', [])
                    if isinstance(descriptions, str):
                        descriptions = [descriptions]
                    
                    for desc in descriptions:
                        if desc.strip():
                            bullet_para = doc.add_paragraph(desc.strip(), style='List Bullet')
                            bullet_para.runs[0].font.size = Pt(10)
        
        # Certifications
        if resume_data.get('certifications'):
            self._add_section_heading(doc, 'CERTIFICATIONS')
            certs = self._parse_json_field(resume_data['certifications'])
            
            for cert in certs:
                if isinstance(cert, dict):
                    cert_para = doc.add_paragraph(cert.get('title', 'Certification'), style='List Bullet')
                    cert_para.runs[0].font.size = Pt(10)
                elif isinstance(cert, str):
                    cert_para = doc.add_paragraph(cert, style='List Bullet')
                    cert_para.runs[0].font.size = Pt(10)
        
        # Save document
        doc.save(str(filepath))
        return str(filepath)
    
    def _add_section_heading(self, doc, heading_text):
        """Add a formatted section heading to the document."""
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(heading_text)
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.color.rgb = RGBColor(44, 90, 160)  # Professional blue
        
        # Add a bottom border
        p_pr = heading_para._element.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2c5aa0')
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
    
    def _parse_json_field(self, field):
        """Parse a field that might be JSON string or already parsed."""
        if isinstance(field, str):
            try:
                return json.loads(field)
            except:
                return field
        return field
    
    def generate_txt(self, resume_data: Dict, filename: Optional[str] = None) -> str:
        """Generate a plain text resume (ATS-friendly)."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.txt"
        
        filepath = self.output_dir / filename
        
        lines = []
        
        # Name
        name = resume_data.get('name', resume_data.get('full_name', 'Your Name'))
        lines.append(name.upper())
        lines.append("=" * len(name))
        lines.append("")
        
        # Contact
        contact_parts = []
        if resume_data.get('email'):
            contact_parts.append(f"Email: {resume_data['email']}")
        if resume_data.get('phone'):
            contact_parts.append(f"Phone: {resume_data['phone']}")
        if resume_data.get('location'):
            contact_parts.append(f"Location: {resume_data['location']}")
        
        lines.extend(contact_parts)
        lines.append("")
        lines.append("")
        
        # Summary
        if resume_data.get('summary'):
            lines.append("PROFESSIONAL SUMMARY")
            lines.append("-" * 40)
            lines.append(resume_data['summary'])
            lines.append("")
            lines.append("")
        
        # Skills
        if resume_data.get('skills'):
            lines.append("SKILLS")
            lines.append("-" * 40)
            skills = self._parse_json_field(resume_data['skills'])
            if isinstance(skills, list):
                lines.append(", ".join(skills))
            else:
                lines.append(str(skills))
            lines.append("")
            lines.append("")
        
        # Experience
        if resume_data.get('experience'):
            lines.append("PROFESSIONAL EXPERIENCE")
            lines.append("-" * 40)
            experiences = self._parse_json_field(resume_data['experience'])
            
            for exp in experiences:
                if isinstance(exp, dict):
                    lines.append(exp.get('title', 'Position'))
                    lines.append(exp.get('company', ''))
                    lines.append("")
                    
                    descriptions = exp.get('description', [])
                    if isinstance(descriptions, str):
                        descriptions = [descriptions]
                    
                    for desc in descriptions:
                        if desc.strip():
                            lines.append(f"  • {desc.strip()}")
                    
                    lines.append("")
            lines.append("")
        
        # Education
        if resume_data.get('education'):
            lines.append("EDUCATION")
            lines.append("-" * 40)
            educations = self._parse_json_field(resume_data['education'])
            
            for edu in educations:
                if isinstance(edu, dict):
                    lines.append(edu.get('title', 'Degree'))
                    
                    descriptions = edu.get('description', [])
                    if isinstance(descriptions, str):
                        descriptions = [descriptions]
                    
                    for desc in descriptions:
                        if desc.strip():
                            lines.append(f"  {desc.strip()}")
                    
                    lines.append("")
        
        # Write to file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
        
        return str(filepath)

